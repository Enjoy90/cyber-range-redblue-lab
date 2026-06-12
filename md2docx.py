#!/usr/bin/env python3
# =====================================================================
#  md2docx.py  -  Rakit SUBMISSION.md menjadi 1 .docx PROFESIONAL
# ---------------------------------------------------------------------
#  - Gaya monokrom/korporat (tanpa warna mencolok), tipografi bersih.
#  - Lampiran A (source code) di-embed OTOMATIS dari berkas asli,
#    sehingga dokumen selalu sinkron dengan kode terbaru.
#  Output: Cyber_Range_Submission.docx
# =====================================================================
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Cyber_Range_Submission.docx")

# Palet monokrom (tanpa warna "AI")
INK      = RGBColor(0x1A, 0x1A, 0x1A)   # teks utama
HEAD     = RGBColor(0x22, 0x2A, 0x33)   # heading (slate gelap, bukan biru)
CODECLR  = RGBColor(0x2B, 0x2B, 0x2B)   # teks kode
INLINECLR= RGBColor(0x33, 0x33, 0x33)   # inline code
RULECLR  = "AAAAAA"
HDRFILL  = "E8E8E8"                     # shading header tabel
CODEFILL = "F4F4F4"                     # shading blok kode

# Berkas sumber yang di-embed otomatis di Lampiran A (urut logis)
SOURCES = [
    "app/package.json", "app/server.js", "app/Dockerfile",
    "nginx/default.conf",
    "admin-bot/package.json", "admin-bot/bot.js", "admin-bot/Dockerfile",
    "docker-compose.yml",
    "scripts/inject_logs.py", "scripts/bootstrap.sh", "scripts/reset.sh",
    "cloud-init/user-data",
]

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

INLINE = re.compile(r"(`[^`]*`|\*\*[^*]+\*\*)")


# ---------- helper low-level (shading & border via OXML) -------------
def _shade(pr, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pr.append(shd)

def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)

def shade_para(p, fill):
    _shade(p._p.get_or_add_pPr(), fill)

def bottom_border(p, sz=6, color=RULECLR):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '6'); b.set(qn('w:color'), color)
    pbdr.append(b); pPr.append(pbdr)


# ---------- rendering markdown ---------------------------------------
def add_inline(par, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1]); r.font.name = "Consolas"
            r.font.size = Pt(9.5); r.font.color.rgb = INLINECLR
        elif part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
        else:
            par.add_run(part)

def add_heading(level, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    sizes = {1: 19, 2: 14.5, 3: 12, 4: 11}
    add_inline(p, text)
    for r in p.runs:
        r.bold = True; r.font.color.rgb = HEAD
        r.font.size = Pt(sizes.get(level, 11)); r.font.name = "Calibri"
    if level <= 2:
        bottom_border(p)
    return p

def add_code_block(lines):
    for ln in (lines or [" "]):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
        pf.left_indent = Inches(0.12); pf.right_indent = Inches(0.05)
        shade_para(p, CODEFILL)
        r = p.add_run(ln if ln != "" else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = CODECLR

def split_row(line):
    s = line.strip()
    if s.startswith("|"): s = s[1:]
    if s.endswith("|"): s = s[:-1]
    return s.split("|")

def add_table(rows):
    header = rows[0]
    body = rows[2:] if len(rows) > 2 else []
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, c in enumerate(header):
        cell = t.rows[0].cells[i]
        shade_cell(cell, HDRFILL)
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], c.strip())
        for r in cell.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(9.5)
    for row in body:
        cells = t.add_row().cells
        for i in range(len(header)):
            val = row[i].strip() if i < len(row) else ""
            cells[i].paragraphs[0].text = ""
            add_inline(cells[i].paragraphs[0], val)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def render(md):
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # lewati komentar HTML <!-- ... -->
        if line.lstrip().startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1; continue

        # code block
        if line.lstrip().startswith("```"):
            i += 1; buf = []
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i]); i += 1
            add_code_block(buf); i += 1; continue

        # tabel
        if line.strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|", lines[i+1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i])); i += 1
            add_table(rows); continue

        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            add_heading(len(m.group(1)), m.group(2).strip()); i += 1; continue

        if re.match(r"^\s*---+\s*$", line):
            i += 1; continue

        if line.strip().startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            add_inline(p, re.sub(r"^\s*>\s?", "", line))
            for r in p.runs:
                r.italic = True; r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1; continue

        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^\s*[-*]\s+", "", line)); i += 1; continue

        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\s*\d+\.\s+", "", line)); i += 1; continue

        if line.strip() == "":
            i += 1; continue

        p = doc.add_paragraph(); add_inline(p, line); i += 1


# ---------- footer: nomor halaman ------------------------------------
def add_page_numbers():
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Cyber Range Engineering — Red vs. Blue Lab    |    ").font.size = Pt(8)
    run = p.add_run()
    for kind, val in (("begin", None), ("instr", "PAGE"), ("end", None)):
        if kind == "instr":
            el = OxmlElement('w:instrText'); el.set(qn('xml:space'), 'preserve'); el.text = val
        else:
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), kind if kind != "begin" else "begin")
        run._r.append(el)
    run.font.size = Pt(8)


# ---------- build ----------------------------------------------------
with open(os.path.join(HERE, "SUBMISSION.md"), encoding="utf-8") as f:
    render(f.read())

# Lampiran A: embed seluruh source code dari berkas ASLI
doc.add_page_break()
add_heading(2, "Lampiran A — Source Code Lengkap")
p = doc.add_paragraph()
add_inline(p, "Salinan utuh seluruh berkas sumber repository (dibaca otomatis dari kode "
              "terbaru saat dokumen ini dibuat), agar dokumen bersifat self-contained.")
for idx, rel in enumerate(SOURCES, 1):
    path = os.path.join(HERE, rel.replace("/", os.sep))
    add_heading(3, f"A.{idx}  {rel}")
    try:
        with open(path, encoding="utf-8") as fh:
            add_code_block(fh.read().split("\n"))
    except Exception as e:
        doc.add_paragraph(f"[tidak dapat membaca {rel}: {e}]")

add_page_numbers()
doc.save(OUT)
print("OK ->", OUT)
